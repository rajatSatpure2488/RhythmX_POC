"""Generate a Word document describing the exact DrChrono JSON payloads MediSync
pushes for 8 resources. Payloads mirror the mappers in app/routes/push.py."""
import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join("docs", "DrChrono_Push_JSON_Structures.docx")

# (Title, METHOD, endpoint, content-type, notes, payload dict OR None for multipart, multipart_fields)
RESOURCES = [
    (
        "1. Patient Creation", "POST", "/api/patients", "application/json",
        "Creates a patient. Sent as flat JSON. `doctor` is the DrChrono provider ID. "
        "`gender` must be Male / Female / Other. `date_of_birth` must be YYYY-MM-DD.",
        {
            "first_name": "Samuel",
            "last_name": "Rossi",
            "date_of_birth": "1945-01-15",
            "gender": "Male",
            "email": "samuel.rossi@example.com",
            "home_phone": "(512) 555-0142",
            "address": "123 Main St",
            "city": "Austin",
            "state": "TX",
            "zip_code": "78701",
            "doctor": 525460,
        },
        None,
    ),
    (
        "2. Medication Creation", "POST", "/api/medications", "application/json",
        "Creates a medication on the patient's chart. `status` is lowercase "
        "(active / inactive). `rxnorm` and dosage fields are optional.",
        {
            "patient": 134185516,
            "doctor": 525460,
            "name": "Amlodipine 5 MG Oral Tablet",
            "status": "active",
            "rxnorm": "197361",
            "frequency": "1 tablet once daily",
            "route": "oral",
            "indication": "Hypertension",
            "start_date": "2024-07-01",
        },
        None,
    ),
    (
        "3. Allergies Creation", "POST", "/api/allergies", "application/json",
        "Creates an allergy/intolerance. `description` is the allergen. "
        "`status` is lowercase (active / inactive).",
        {
            "patient": 134185516,
            "doctor": 525460,
            "description": "Penicillin",
            "status": "active",
            "reaction": "Hives",
            "notes": "Patient reports rash on exposure.",
        },
        None,
    ),
    (
        "4. Problem Creation", "POST", "/api/problems", "application/json",
        "Creates a problem/condition. `description` is required. `icd_code` is the "
        "ICD-10 code. `status` is lowercase (active / resolved).",
        {
            "patient": 134185516,
            "doctor": 525460,
            "description": "Essential (primary) hypertension",
            "status": "active",
            "icd_code": "I10",
            "date_onset": "1997-03-31",
        },
        None,
    ),
    (
        "5. Appointment Creation", "POST", "/api/appointments", "application/json",
        "Creates an appointment. `scheduled_time` is ISO-8601 (YYYY-MM-DDTHH:MM:SS, "
        "year 2000-2099 only). `duration` is minutes. `office` + `exam_room` are "
        "required by DrChrono; `office` is auto-resolved from the doctor when absent. "
        "`status` must be a valid enum (Confirmed / Complete / Cancelled / ...). "
        "`reason` is capped at 100 characters.",
        {
            "patient": 134185516,
            "doctor": 525460,
            "scheduled_time": "2024-07-01T10:00:00",
            "duration": 30,
            "status": "Confirmed",
            "reason": "Follow-up for hypertension management",
            "allow_overlapping": True,
            "office": 559437,
            "exam_room": 1,
        },
        None,
    ),
    (
        "6. Encounter Creation", "POST", "/api/appointments", "application/json",
        "DrChrono has NO encounters endpoint — each encounter is pushed as an "
        "appointment using the SAME structure as #5. `scheduled_time` comes from the "
        "encounter start; `duration` from start/end; `notes` carries encounter type "
        "and specialty.",
        {
            "patient": 134185516,
            "doctor": 525460,
            "scheduled_time": "2024-07-01T10:00:00",
            "duration": 75,
            "status": "Complete",
            "reason": "Office Visit - Cardiology",
            "allow_overlapping": True,
            "office": 559437,
            "exam_room": 1,
        },
        None,
    ),
    (
        "7. Diagnostic Reports", "POST", "/api/documents", "multipart/form-data",
        "DrChrono's lab API is partner-gated (403), so each diagnostic report is "
        "rendered to a PDF and uploaded as a clinical Document. Sent as multipart "
        "form-data (NOT JSON): text fields + a generated PDF file. `metatags` is a "
        "JSON-array string.",
        None,
        [
            ("patient", "134185516", "form field"),
            ("doctor", "525460", "form field"),
            ("description", "Laboratory", "form field (<=100 chars)"),
            ("date", "2024-07-01", "form field (YYYY-MM-DD)"),
            ("metatags", '["diagnostic_report", "Laboratory"]', "form field (JSON array string)"),
            ("document", "diagnostic_report_<id>.pdf  (binary PDF)", "file part, mime application/pdf"),
        ],
    ),
    (
        "8. Documents Upload", "POST", "/api/documents", "multipart/form-data",
        "Direct clinical document/file upload. Sent as multipart form-data: text "
        "fields + the binary file (PDF/JPG/PNG/TIFF). `metatags` is pipe-separated; "
        "`archived` is optional.",
        None,
        [
            ("patient", "134185516", "form field"),
            ("doctor", "525460", "form field"),
            ("description", "Discharge Summary", "form field"),
            ("date", "2024-07-01", "form field (YYYY-MM-DD)"),
            ("metatags", "lab|cbc|uploaded", "form field (pipe-separated)"),
            ("archived", "false", "form field (optional)"),
            ("document", "summary.pdf  (binary file)", "file part, mime application/pdf"),
        ],
    ),
    (
        "9. Coverage / Insurance", "POST", "/api/insurances", "application/json",
        "Patient insurance. There is NO /api/patient_insurances endpoint (it 404s). "
        "Primary vs secondary is conveyed by `insurance_type`.",
        {
            "patient": 134185516,
            "insurance_type": "primary",
            "insurance_company": "UnitedHealthcare",
            "insurance_plan_name": "Choice Plus PPO 1500",
            "insurance_id_number": "98765432101",
            "insurance_group_number": "7890012",
            "payer_id": "789",
        },
        None,
    ),
    (
        "10. Clinical Note — Vitals (Step 1 of 2)", "PATCH",
        "/api/appointments/{appointment_id}", "application/json",
        "Clinical notes push in TWO steps. STEP 1 writes vitals onto the appointment "
        "(success = 204 No Content). Units are fixed: temperature='f', "
        "height/head='inches', weight='lbs'. Only vitals actually present are sent; "
        "temperature is converted Celsius->Fahrenheit.",
        {
            "vitals": {
                "temperature": 98.2, "temperature_units": "f",
                "height": 68, "height_units": "inches",
                "weight": 174, "weight_units": "lbs",
                "head_circumference": 0, "head_circumference_units": "inches",
                "blood_pressure_1": 160, "blood_pressure_2": 90,
                "pulse": 82, "respiratory_rate": 16, "oxygen_saturation": 98,
                "pain": "1", "smoking_status": "blank",
            },
            "status": "Checked In",
        },
        None,
    ),
    (
        "11. Clinical Note — Note (Step 2 of 2)", "POST",
        "/api/yellow_notepad?appointment_id={id}&template_id=7520906", "application/json",
        "STEP 2 posts the clinical narrative (success = 200/201). appointment_id and "
        "template_id (fixed 7520906) are QUERY PARAMETERS; the body holds only `content` "
        "(CC/HPI/Assessment/Plan — vitals are NOT repeated here, they live on the "
        "appointment). appointment_id is resolved from the appointment pushed in the run "
        "or a live DrChrono lookup.",
        {
            "content": "CC: HTN follow-up  HPI: 68yo male, headaches.  "
                       "Assessment: Essential hypertension  Plan: Increase amlodipine.",
        },
        None,
    ),
    (
        "12. Observations + Observation Notes", "POST", "/api/patient_lab_results",
        "application/json",
        "Observations are pushed as structured lab results (one per observation), "
        "enriched with their matching observation note (LEFT join on observation_id). "
        "Key value rules learned from DrChrono: date_test_performed must be a full "
        "ISO-8601 datetime (NOT date-only); lab_order_status valid choices are "
        "'In Progress' / 'Reviewed' (NOT 'Resulted'); loinc_code only when code_vocab "
        "is LOINC; missing fields default to 'Not provided'.",
        {
            "ordering_doctor": 525460,
            "patient": 134185516,
            "title": "Sodium [Moles/volume] in Serum or Plasma",
            "loinc_code": "2951-2",
            "lab_result_value": "140 mEq/L. Imported via RhythmX AI integration pipeline.",
            "lab_result_value_as_float": 140.0,
            "lab_result_value_units": "mEq/L",
            "lab_normal_range": "135-145 mEq/L",
            "lab_normal_range_units": "mEq/L",
            "lab_abnormal_flag": "",
            "lab_order_status": "Reviewed",
            "date_test_performed": "1997-03-31T00:00:00",
            "doctor_signoff": False,
            "doctor_comments": "Observation Note: Within normal limits. Category: laboratory "
                               "Result imported through RhythmX AI API integration workflow.",
        },
        None,
    ),
]

doc = Document()

# ── Styles ──
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

title = doc.add_heading("DrChrono Push — JSON / Payload Structures", level=0)
sub = doc.add_paragraph(
    "Exact request structures MediSync sends to DrChrono for each resource. "
    "Payloads are flat DrChrono JSON (not FHIR). All requests carry the header "
    "Authorization: Bearer <access_token> and X-DRC-API-Version: v4. "
    "Updated to include Coverage/Insurance, the two-step Clinical Note push "
    "(appointment vitals + yellow_notepad), and Observations -> patient_lab_results."
)
sub.runs[0].italic = True

note = doc.add_paragraph()
note.add_run("Reference — value rules confirmed from live DrChrono responses:").bold = True
for line in [
    "Patient gender: Male / Female / Other only ('Unknown' is rejected); read from gender_administrative.",
    "Dates: YYYY-MM-DD; date_test_performed needs a full ISO datetime (YYYY-MM-DDThh:mm:ss).",
    "Appointment status: a fixed enum (Confirmed/Complete/Cancelled/...); year must be 2000-2099.",
    "Lab order status: 'In Progress' / 'Reviewed' (NOT 'Resulted').",
    "Insurance: POST /api/insurances with insurance_type primary/secondary (no /api/patient_insurances).",
    "Diagnostic reports / clinical notes / observations that can't use a gated API are sent as the shapes shown here.",
]:
    doc.add_paragraph(line, style="List Bullet")
doc.add_paragraph("")


def add_json_block(payload: dict):
    text = json.dumps(payload, indent=2)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_table(fields):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Field", "Example value", "Type / notes"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
    for name, val, note in fields:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = name, val, note


for entry in RESOURCES:
    name, method, endpoint, ctype, notes, payload, multipart = entry
    doc.add_heading(name, level=1)

    meta = doc.add_paragraph()
    meta.add_run("Endpoint: ").bold = True
    meta.add_run(f"{method} {endpoint}\n")
    meta.add_run("Content-Type: ").bold = True
    meta.add_run(f"{ctype}\n")
    meta.add_run("Notes: ").bold = True
    meta.add_run(notes)

    if payload is not None:
        doc.add_paragraph("Request body (JSON):").runs[0].bold = True
        add_json_block(payload)
    else:
        doc.add_paragraph("Request body (multipart/form-data fields):").runs[0].bold = True
        add_table(multipart)

    doc.add_paragraph("")

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Resources documented: {len(RESOURCES)}")
